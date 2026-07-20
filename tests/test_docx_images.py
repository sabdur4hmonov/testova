"""
DOCX embedded-image extraction → per-question image_path → rendered variant.

Regression for the bug where a DOCX with real embedded pictures (word/media/
image*.png) produced only a "[Rasm]: ..." text description in the generated
variant instead of the actual image. See attach_docx_inline_images (extraction)
and _load_image_rl's docximg_ sizing branch (rendering).
"""
import io

import fitz  # PyMuPDF
import pytest
from docx import Document
from PIL import Image

from app.services.file_processor import (
    attach_docx_inline_images,
    docx_extract_ordered_images,
)
from app.services.pdf_generator import (
    build_variants_pdf,
    build_variants_pdf_compact,
)


# ── helpers ───────────────────────────────────────────────────────────────────

def _green_png(color=(0, 128, 0)) -> io.BytesIO:
    img = Image.new("RGB", (400, 250), color)
    b = io.BytesIO()
    img.save(b, "PNG")
    b.seek(0)
    return b


def _docx_with_inline_images(n: int) -> bytes:
    """A test where questions 1..n each carry one inline picture."""
    doc = Document()
    for i in range(1, n + 1):
        doc.add_paragraph(f"{i}. Rasmga qarang, savol {i}?")
        doc.add_picture(_green_png())
        doc.add_paragraph("A) x  B) y  C) z  D) w")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _flagged_questions(n: int) -> list[dict]:
    return [
        {
            "question_number": i,
            "position_in_variant": i,
            "question_text": f"Rasmga qarang, savol {i}?",
            "options": {"A": "x", "B": "y", "C": "z", "D": "w"},
            "has_image": True,
            "image_path": None,
            "image_description": "A stylized green tree on a light blue background.",
        }
        for i in range(1, n + 1)
    ]


def _variant(questions: list[dict]) -> dict:
    return {"variant_number": 1, "questions_data": questions}


def _pdf_text_and_img_count(pdf_bytes: bytes) -> tuple[str, int]:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = "".join(p.get_text() for p in doc)
    imgs = sum(len(p.get_images()) for p in doc)
    doc.close()
    return text, imgs


# ── extraction ────────────────────────────────────────────────────────────────

def test_docx_ordered_images_found():
    imgs = docx_extract_ordered_images(_docx_with_inline_images(3))
    assert len(imgs) == 3
    assert all(isinstance(p, Image.Image) for p in imgs)


def test_attach_populates_image_path_count_matched():
    docx = _docx_with_inline_images(3)
    questions = _flagged_questions(3)
    attached = attach_docx_inline_images(questions, docx)
    assert attached == 3
    for q in questions:
        assert q["image_path"], "flagged question should get an image_path"
        # the docximg_ marker is what tells the renderer to fit-to-column
        assert "docximg_" in q["image_path"]


def test_attach_count_mismatch_attaches_nothing():
    # 3 images in the doc, but only 2 flagged questions → never guess.
    docx = _docx_with_inline_images(3)
    questions = _flagged_questions(2)
    attached = attach_docx_inline_images(questions, docx)
    assert attached == 0
    assert all(q["image_path"] is None for q in questions)


def test_attach_skips_questions_without_has_image():
    docx = _docx_with_inline_images(1)
    questions = _flagged_questions(1)
    questions[0]["has_image"] = False
    # flagged count (0) != images (1) → nothing attached
    assert attach_docx_inline_images(questions, docx) == 0
    assert questions[0]["image_path"] is None


def test_attach_bad_docx_bytes_no_crash():
    questions = _flagged_questions(1)
    assert attach_docx_inline_images(questions, b"not a real docx") == 0
    assert questions[0]["image_path"] is None


# ── rendering: image_path wins over description (both layouts) ─────────────────

@pytest.mark.parametrize("build", [build_variants_pdf, build_variants_pdf_compact])
def test_image_path_renders_image_not_description(build):
    docx = _docx_with_inline_images(1)
    questions = _flagged_questions(1)
    attach_docx_inline_images(questions, docx)
    assert questions[0]["image_path"]

    pdf = build([_variant(questions)])
    text, img_count = _pdf_text_and_img_count(pdf)
    assert img_count >= 1, "the embedded image must be drawn"
    assert "[Rasm]" not in text, "description box must NOT appear when the image renders"


# ── rendering: description-only still shows the description (both layouts) ─────

@pytest.mark.parametrize("build", [build_variants_pdf, build_variants_pdf_compact])
def test_description_only_shows_description(build):
    q = _flagged_questions(1)[0]
    q["image_path"] = None  # no image, only a description
    pdf = build([_variant([q])])
    text, img_count = _pdf_text_and_img_count(pdf)
    assert "[Rasm]" in text, "with no image_path the description box must render"


# ── rendering: missing file falls back gracefully (both layouts) ──────────────

@pytest.mark.parametrize("build", [build_variants_pdf, build_variants_pdf_compact])
def test_missing_image_file_falls_back_to_description(build):
    q = _flagged_questions(1)[0]
    q["image_path"] = "temp_images/docximg_DOES_NOT_EXIST.png"
    # must not raise, and must fall back to the description
    pdf = build([_variant([q])])
    text, _ = _pdf_text_and_img_count(pdf)
    assert "[Rasm]" in text, "a missing image file must fall back to the description"
