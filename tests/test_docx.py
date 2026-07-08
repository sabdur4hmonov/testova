"""
DOCX rendering — a typed test (text + options, no embedded pictures) must
produce real page images for the vision pipeline, not zero pages.
"""
import io

import numpy as np
from docx import Document

from app.services.file_processor import docx_to_images


def _build_docx(paragraphs, table_rows=None) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                t.rows[r].cells[c].text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _ink_fraction(pil_img) -> float:
    arr = np.asarray(pil_img.convert("L"))
    return float((arr < 128).mean())


THREE_QUESTIONS = [
    "1. Suvning kimyoviy formulasi qanday?",
    "A) H2O  B) CO2  C) O2  D) H2",
    "2. Eng yengil metall qaysi?",
    "A) Temir  B) Litiy  C) Mis  D) Rux",
    "3. 2 + 2 necha bo'ladi?",
    "A) 3  B) 4  C) 5  D) 22",
]


def test_typed_docx_produces_pages():
    pages, text = docx_to_images(_build_docx(THREE_QUESTIONS))
    assert len(pages) >= 1, "a typed DOCX must produce at least one page image"
    # The page must actually contain rendered ink (not a blank canvas)
    assert _ink_fraction(pages[0].image) > 0.002
    # Text content carries every question + options for downstream sanity
    for q in THREE_QUESTIONS:
        assert q.split(")")[0][:6] in text or q[:6] in text


def test_docx_text_has_all_questions_and_options():
    _pages, text = docx_to_images(_build_docx(THREE_QUESTIONS))
    for marker in ("1.", "2.", "3.", "H2O", "Litiy"):
        assert marker in text


def test_docx_with_table_renders():
    pages, text = docx_to_images(_build_docx(
        ["1. Jadvalga qarang:"],
        table_rows=[["Modda", "Formula"], ["Suv", "H2O"], ["Tuz", "NaCl"]],
    ))
    assert len(pages) >= 1
    assert "H2O" in text and "NaCl" in text
    assert _ink_fraction(pages[0].image) > 0.002


def test_long_docx_paginates():
    many = [f"{i}. Savol raqami {i} — matn shu yerda davom etadi." for i in range(1, 80)]
    pages, _text = docx_to_images(_build_docx(many))
    assert len(pages) >= 2, "80 questions must overflow onto multiple pages"


def test_page_dimensions_match_pipeline_dpi():
    from app.services.file_processor import _DOCX_PAGE_W, _DOCX_PAGE_H
    pages, _ = docx_to_images(_build_docx(["1. Test"]))
    assert pages[0].image.size == (_DOCX_PAGE_W, _DOCX_PAGE_H)


def test_empty_docx_does_not_crash():
    pages, text = docx_to_images(_build_docx([]))
    assert len(pages) >= 1  # at least a blank page, never an exception
    assert text == ""
