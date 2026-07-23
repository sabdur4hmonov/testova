"""
DOCX superscript/subscript recovery (Defect 1). python-docx `.text` drops
run-level `w:vertAlign`, flattening U²→U2, 2⁸→28 (arithmetically false). We
rebuild paragraph text from runs and surface the formatting as ^/_ notation so
Gemini transcribes 2^8 and math_render typesets 2⁸.

Includes: the trailing-separator case found in the real paper (a ';' inheriting
superscript must NOT be swallowed into the exponent), adversarial cases proving
correct render, and the Defect 2 regression tests (scripts bind tighter than
fraction division: a/b^2 is a/b², not the different number (a/b)²).
"""
import io

from docx import Document

from app.services.file_processor import (
    _wrap_script, _para_scripted_text, docx_to_images,
)
from app.services.math_render import parse


# ── _wrap_script units (incl. the real trailing-separator finding) ───────────
def test_wrap_single_char():
    assert _wrap_script("2", "^") == "^2"
    assert _wrap_script("1", "_") == "_1"


def test_wrap_multidigit_no_parens():
    assert _wrap_script("10", "^") == "^10"      # bare number, parses whole


def test_wrap_multitoken_gets_parens():
    assert _wrap_script("n+1", "^") == "^(n+1)"


def test_wrap_peels_trailing_separator():
    # the observed paper: superscript run was "8;  " / "10;" — the ';' separator
    # must stay OUT of the exponent
    assert _wrap_script("8;  ", "^") == "^8;  "
    assert _wrap_script("10;", "^") == "^10;"


def test_wrap_only_punctuation_left_alone():
    assert _wrap_script(";", "^") == ";"


# ── DOCX round-trip: runs → ^/_ text ─────────────────────────────────────────
def _docx(runs):
    """runs: list of (text, script) where script in {None,'sup','sub'}."""
    doc = Document()
    p = doc.add_paragraph()
    for text, script in runs:
        r = p.add_run(text)
        if script == "sup":
            r.font.superscript = True
        elif script == "sub":
            r.font.subscript = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_superscript_becomes_caret():
    _pages, text = docx_to_images(_docx([("NUK=U", None), ("2", "sup"), (".", None)]))
    assert "U^2" in text and "U2" not in text.replace("U^2", "")


def test_docx_subscript_becomes_underscore():
    _pages, text = docx_to_images(_docx([("x", None), ("1", "sub")]))
    assert "x_1" in text


def test_docx_option_separator_not_swallowed():
    # "512=2" + superscript "8;  " → must be "512=2^8;" not "512=2^(8;)"
    _pages, text = docx_to_images(_docx([("512=2", None), ("8;  ", "sup")]))
    assert "512=2^8" in text
    assert "2^(8;" not in text and "2^8;" in text.replace("  ", "")


def test_docx_plain_text_unchanged():
    _pages, text = docx_to_images(_docx([("Suvning formulasi?", None)]))
    assert "Suvning formulasi?" in text and "^" not in text


# ── adversarial: the emitted notation renders CORRECTLY ──────────────────────
def test_observed_exponents_render_correct():
    assert parse("128=2^7").latex() == "128={2}^{7}"
    assert parse("1024=2^10").latex() == "1024={2}^{10}"    # whole exponent
    assert parse("143 sm^2").latex() == "143 {sm}^{2}"


def test_paren_base_and_multitoken_exponent_correct():
    assert parse("(a+b)^2").latex() == "{\\left(a+b\\right)}^{2}"
    # multi-token exponent next to a paren stem — no bleed
    assert parse("2^(n+1) + (c-1)").latex() == "{2}^{n+1} + \\left(c-1\\right)"


# ── Defect 2 (FIXED): scripts bind tighter than fraction division ────────────
def test_denominator_script_binds_tighter_than_fraction():
    # A script on a denominator belongs to the DENOMINATOR, never hoisted onto
    # the whole fraction — (a/b)^2 is a DIFFERENT NUMBER than a/b^2, so the old
    # hoist put false math in front of students.
    assert parse("a/b^2").latex() == "\\frac{a}{{b}^{2}}"
    assert parse("a^2/b^2").latex() == "\\frac{{a}^{2}}{{b}^{2}}"
    assert parse("a/b^2/c").latex() == "\\frac{\\frac{a}{{b}^{2}}}{c}"
    # a radical numerator stays intact; only the precedence around it is fixed
    assert parse("sqrt(a)/b^2").latex() == "\\frac{\\sqrt{a}}{{b}^{2}}"
    # a parenthesised denominator takes the script too
    assert parse("a/(b+c)^2").latex() == "\\frac{a}{{\\left(b+c\\right)}^{2}}"
    # subscripts are the SAME rule: x_1/x_2 is x1/x2, never (x_1/x)_2
    assert parse("a/b_2").latex() == "\\frac{a}{{b}_{2}}"
    assert parse("x_1/x_2").latex() == "\\frac{{x}_{1}}{{x}_{2}}"


def test_defect2_fix_leaves_neighbouring_shapes_alone():
    # the explicit escape hatch still raises the WHOLE fraction
    assert parse("(a/b)^2").latex() == "{\\left(\\frac{a}{b}\\right)}^{2}"
    # a numerator script was always correct and must not move
    assert parse("a^2/b").latex() == "\\frac{{a}^{2}}{b}"
    # division stays left-associative
    assert parse("a/b/c").latex() == "\\frac{\\frac{a}{b}}{c}"
    # a radical denominator is untouched
    assert parse("a/sqrt(b)").latex() == "\\frac{a}{\\sqrt{b}}"
