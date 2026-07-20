"""Ingest uploaded files and prepare them for OCR/AI analysis."""
from __future__ import annotations

import io
import re
import uuid
from pathlib import Path
from typing import NamedTuple

import fitz  # PyMuPDF
from docx import Document
from docx.document import Document as _DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as _DocxTable
from docx.text.paragraph import Paragraph as _DocxParagraph
from PIL import Image, ImageDraw, ImageFont

from app.utils.logging import get_logger

logger = get_logger(__name__)

SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "image/jpeg": "image",
    "image/png": "image",
    "image/webp": "image",
}

# Where cropped question images are saved
IMAGE_SAVE_DIR = Path("temp_images")
DPI = 200  # must match pdf_to_images DPI
# Figure crops are re-rendered straight from the PDF at this DPI so they are
# sharp in the generated variants (the 200-DPI page render was blurry).
# pdf_generator sizes images assuming this density — keep them in sync.
CROP_DPI = 400

# ── BUG FIX: Max ratio of image area vs page area.
# If an embedded image covers >60% of the page, it's a scanned/watermarked
# background page — NOT a question diagram. Skip it.
MAX_IMAGE_PAGE_AREA_RATIO = 0.60


def _ensure_image_dir() -> Path:
    IMAGE_SAVE_DIR.mkdir(parents=True, exist_ok=True)
    return IMAGE_SAVE_DIR


class PageImage(NamedTuple):
    page_number: int
    image: Image.Image


# ── File type detection ───────────────────────────────────────────────────────

def detect_file_type(filename: str, content: bytes) -> str:
    if content[:4] == b"%PDF":
        return "pdf"
    if content[:4] == b"PK\x03\x04":
        return "docx"
    if content[:3] in (b"\xff\xd8\xff", b"\x89PNG", b"GIF"):
        return "image"
    ext = Path(filename).suffix.lower()
    return {"pdf": "pdf", ".pdf": "pdf", ".docx": "docx", ".jpg": "image",
            ".jpeg": "image", ".png": "image"}.get(ext, "unknown")


# ── PDF utilities ─────────────────────────────────────────────────────────────

def pdf_extract_pages_text(pdf_bytes: bytes) -> list[str]:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    result: list[str] = []
    for page_num, page in enumerate(doc, start=1):
        blocks = page.get_text("blocks", sort=True)
        lines: list[str] = []
        for block in blocks:
            if block[6] == 0:
                text = block[4].strip()
                if text:
                    lines.append(text)
        if lines:
            result.append(f"=== Page {page_num} ===\n" + "\n".join(lines))
    doc.close()
    return result


def pdf_to_images(pdf_bytes: bytes, dpi: int = DPI) -> list[PageImage]:
    """Convert every PDF page to a PIL Image.

    Cost optimization: a page with NO embedded raster image renders fine for
    Gemini at a lower DPI (150) — fewer image tiles, cheaper call. Pages that
    carry an embedded image keep the full DPI so figure detail survives. Figure
    CROPS are re-rendered straight from the PDF at CROP_DPI, so the page DPI
    never affects crop sharpness.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages: list[PageImage] = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        page_dpi = dpi if page.get_images() else min(dpi, 150)
        zoom = page_dpi / 72
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        pages.append(PageImage(page_number=page_num + 1, image=img))
        logger.debug("converted_page", page=page_num + 1, total=len(doc), dpi=page_dpi)
    doc.close()
    return pages


def compute_page_infos(
    pdf_bytes: bytes,
    page_images: list["PageImage"],
    col_map: dict[int, dict] | None,
) -> list[dict]:
    """
    Per-analysis-image metadata for the cost optimizer: text length and whether
    the source PDF page carries ANY figure (embedded image OR vector drawing).
    Computed per SOURCE page (conservative): a page with a figure is NEVER a
    skip candidate, so a figure-heavy, text-light page (e.g. the number line)
    is always sent to Gemini.
    """
    src_meta: dict[int, dict] = {}
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception:
        return [{"text_len": 999, "has_visual": True} for _ in page_images]

    def meta(src: int) -> dict:
        if src not in src_meta:
            try:
                pg = doc[src - 1]
                has_visual = bool(pg.get_images()) or bool(pg.get_drawings())
                src_meta[src] = {
                    "text_len": len(pg.get_text().strip()),
                    "has_visual": has_visual,
                }
            except Exception:
                src_meta[src] = {"text_len": 999, "has_visual": True}
        return src_meta[src]

    infos: list[dict] = []
    for p in page_images:
        src = (col_map or {}).get(p.page_number, {}).get("src_page", p.page_number)
        infos.append(meta(src))
    doc.close()
    return infos


# ── Two-column detection & splitting ─────────────────────────────────────────

def _detect_column_split(img: Image.Image) -> float | None:
    """
    Detect a two-column page layout. Returns the split position as a fraction
    of page width (0..1), or None for single-column / ambiguous pages.

    Geometry-based, no tuning to any specific document:
    - per-x ink profile = fraction of body rows (header/footer excluded)
      containing dark pixels at that x
    - a GUTTER is a wide near-zero run in the middle zone of the page
    - a DIVIDER LINE is a thin near-solid run in the same zone
    - conservative guards: both halves must carry substantial ink, otherwise
      the page passes through unsplit (today's behavior).
    """
    import numpy as np

    w0, h0 = img.size
    if w0 < 100 or h0 < 100:
        return None
    target_w = 600
    scale = target_w / w0
    small = img.convert("L").resize((target_w, max(50, int(h0 * scale))))
    a = np.asarray(small)
    h, w = a.shape

    # Exclude header/footer bands — titles often span both columns.
    body = a[int(h * 0.12):int(h * 0.92), :]
    ink = body < 160
    total_ink = int(ink.sum())
    if total_ink < body.size * 0.005:  # nearly blank page
        return None

    col_frac = ink.mean(axis=0)  # fraction of body rows inked, per x
    lo, hi = int(w * 0.33), int(w * 0.67)
    zone = col_frac[lo:hi]

    def _runs(mask) -> list[tuple[int, int]]:
        out, start = [], None
        for i, v in enumerate(mask):
            if v and start is None:
                start = i
            elif not v and start is not None:
                out.append((start, i))
                start = None
        if start is not None:
            out.append((start, len(mask)))
        return out

    split_x: float | None = None

    # Gutter: widest near-empty run, at least ~1.5% of page width
    min_gutter = max(3, int(w * 0.015))
    best = None
    for s, e in _runs(zone < 0.02):
        if e - s >= min_gutter and (best is None or e - s > best[1] - best[0]):
            best = (s, e)
    if best:
        split_x = lo + (best[0] + best[1]) / 2
    else:
        # Divider line: thin, near-solid vertical run
        max_line = max(2, int(w * 0.008))
        for s, e in _runs(zone > 0.65):
            if e - s <= max_line:
                split_x = lo + (s + e) / 2
                break

    if split_x is None:
        return None

    # Both halves must hold a substantial share of the ink
    left_ink = int(ink[:, :int(split_x)].sum())
    right_ink = total_ink - left_ink
    if min(left_ink, right_ink) < total_ink * 0.25:
        return None

    return split_x / w


def split_two_column_pages(
    pages: list[PageImage],
) -> tuple[list[PageImage], dict[int, dict]]:
    """
    Split two-column pages into single-column halves, renumbered sequentially
    in reading order (p1-left, p1-right, p2, p3-left, ...). Gemini then only
    ever sees single-column content: column interleaving becomes impossible
    and image-region geometry is computed within the correct column.

    Returns (new_pages, col_map) where col_map maps each NEW page number to
    {"src_page": original page, "x0": left fraction, "x1": right fraction}.
    Single-column pages pass through unchanged (x0=0.0, x1=1.0).
    """
    new_pages: list[PageImage] = []
    col_map: dict[int, dict] = {}

    for p in pages:
        try:
            split = _detect_column_split(p.image)
        except Exception as e:
            logger.warning("column_detect_failed", page=p.page_number, error=str(e))
            split = None

        if split is None:
            n = len(new_pages) + 1
            new_pages.append(PageImage(page_number=n, image=p.image))
            col_map[n] = {"src_page": p.page_number, "x0": 0.0, "x1": 1.0}
            continue

        w, h = p.image.size
        sx = int(w * split)
        for x0f, x1f, half in (
            (0.0, split, p.image.crop((0, 0, sx, h))),
            (split, 1.0, p.image.crop((sx, 0, w, h))),
        ):
            n = len(new_pages) + 1
            new_pages.append(PageImage(page_number=n, image=half))
            col_map[n] = {"src_page": p.page_number, "x0": x0f, "x1": x1f}
        logger.info(
            "page_split_two_columns",
            src_page=p.page_number, split_frac=round(split, 3),
        )

    return new_pages, col_map


def pdf_extract_page_images(pdf_bytes: bytes) -> dict[int, list[bytes]]:
    """
    Extract embedded bitmap images from each page. {page_num: [png_bytes]}

    BUG FIX: Previously this extracted ALL embedded images including full-page
    scanned/watermarked background pages. Now it skips images that cover more
    than MAX_IMAGE_PAGE_AREA_RATIO of the page area — those are backgrounds,
    not question diagrams.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    result: dict[int, list[bytes]] = {}
    for page_num, page in enumerate(doc, start=1):
        page_rect = page.rect
        page_area = page_rect.width * page_rect.height

        imgs: list[bytes] = []
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base = doc.extract_image(xref)
                if not base or not base.get("image"):
                    continue
                w, h = base.get("width", 0), base.get("height", 0)
                if w < 50 or h < 50:
                    continue

                # ── BUG FIX: skip full-page background images ──────────────
                # Get the bbox of this image on the page to check its area ratio
                try:
                    bbox = page.get_image_bbox(img_info)
                    if bbox and not bbox.is_empty and page_area > 0:
                        img_area = bbox.width * bbox.height
                        ratio = img_area / page_area
                        if ratio > MAX_IMAGE_PAGE_AREA_RATIO:
                            logger.info(
                                "skip_background_image",
                                page=page_num,
                                xref=xref,
                                ratio=round(ratio, 3),
                            )
                            continue
                except Exception:
                    # If we can't get bbox, fall back to pixel-size heuristic:
                    # skip if image is larger than typical page resolution at 200dpi
                    # A4 at 200dpi ≈ 1654×2339 px → area ≈ 3.87M px
                    if w * h > 3_000_000:
                        logger.info(
                            "skip_large_image_no_bbox",
                            page=page_num,
                            xref=xref,
                            size=f"{w}x{h}",
                        )
                        continue
                # ── End BUG FIX ────────────────────────────────────────────

                raw = base["image"]
                ext = base.get("ext", "png").lower()
                if ext != "png":
                    pil = Image.open(io.BytesIO(raw)).convert("RGB")
                    buf = io.BytesIO()
                    pil.save(buf, format="PNG")
                    raw = buf.getvalue()
                imgs.append(raw)
            except Exception as e:
                logger.debug("img_extract_skip", error=str(e))
        if imgs:
            result[page_num] = imgs
    doc.close()
    logger.info("pdf_images_extracted", pages_with_imgs=len(result))
    return result


# ── Smart image-region detection ─────────────────────────────────────────────

def _get_image_rects_on_page(pdf_bytes: bytes, page_num: int) -> list[fitz.Rect]:
    """
    Return bounding boxes of all meaningful images on a PDF page.
    Covers both embedded bitmaps AND vector drawings (geometries, diagrams).
    page_num is 1-based.

    BUG FIX: Now also skips full-page background images using area ratio check,
    consistent with pdf_extract_page_images().
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num - 1]
    page_rect = page.rect
    page_area = page_rect.width * page_rect.height
    rects: list[fitz.Rect] = []

    # 1. Embedded bitmap images
    for img_info in page.get_images(full=True):
        xref = img_info[0]
        try:
            base = doc.extract_image(xref)
            w, h = base.get("width", 0), base.get("height", 0)
            if w < 50 or h < 50:
                continue
            bbox = page.get_image_bbox(img_info)
            if bbox and not bbox.is_empty:
                # ── BUG FIX: skip background images ───────────────────────
                if page_area > 0:
                    ratio = (bbox.width * bbox.height) / page_area
                    if ratio > MAX_IMAGE_PAGE_AREA_RATIO:
                        logger.debug(
                            "rect_skip_background",
                            page=page_num,
                            xref=xref,
                            ratio=round(ratio, 3),
                        )
                        continue
                # ── End BUG FIX ────────────────────────────────────────────
                rects.append(bbox)
        except Exception:
            pass

    # 2. Vector drawings — use drawing paths bounding box
    try:
        drawings = page.get_drawings()
        if drawings:
            drawing_rects = [fitz.Rect(d["rect"]) for d in drawings if d.get("rect")]
            clusters = _cluster_rects(drawing_rects, gap=20)
            for cluster_rect in clusters:
                # Only include if reasonably large (not just borders/lines)
                if cluster_rect.width > 40 and cluster_rect.height > 40:
                    rects.append(cluster_rect)
    except Exception as e:
        logger.debug("drawings_skip", error=str(e))

    doc.close()
    return rects


def _cluster_rects(rects: list[fitz.Rect], gap: float = 20) -> list[fitz.Rect]:
    """Merge nearby rects into clusters. Returns one bounding rect per cluster."""
    if not rects:
        return []
    clusters: list[fitz.Rect] = []
    used = [False] * len(rects)
    for i, r in enumerate(rects):
        if used[i]:
            continue
        cluster = fitz.Rect(r)
        used[i] = True
        changed = True
        while changed:
            changed = False
            for j, r2 in enumerate(rects):
                if used[j]:
                    continue
                expanded = fitz.Rect(
                    cluster.x0 - gap, cluster.y0 - gap,
                    cluster.x1 + gap, cluster.y1 + gap,
                )
                if expanded.intersects(r2):
                    cluster |= r2
                    used[j] = True
                    changed = True
        clusters.append(cluster)
    return clusters


def _get_question_y_positions(
    pdf_bytes: bytes,
    page_num: int,
    questions_on_page: list[dict],
    x_range: tuple[float, float] | None = None,
) -> dict[int, float]:
    """
    Get the ACTUAL Y position of each question's text on the page by searching
    for the question-number text in the PDF text blocks.

    x_range (PDF points): when the page is two-column, restrict the search to
    that column's horizontal band — otherwise "12." from the OTHER column
    poisons the geometry (the garbage-crop bug).

    Questions whose number text cannot be located are OMITTED — never
    estimated. Estimated positions produced crops containing other
    questions' content; no image is better than a wrong one.

    Returns: {question_number: y_position_in_pdf_points}
    """
    if not pdf_bytes or not questions_on_page:
        return {}

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num - 1]

    # WORD-level positions, not blocks: PyMuPDF merges same-baseline text
    # from BOTH columns into one block ("1. Left ... 3. Right ..."), which
    # both breaks the startswith match and mis-centers the x filter.
    # Words never merge across columns.
    words = page.get_text("words")  # (x0, y0, x1, y1, text, block, line, word)
    doc.close()

    if x_range is not None:
        words = [
            w for w in words
            if x_range[0] <= (w[0] + w[2]) / 2 <= x_range[1]
        ]

    y_positions: dict[int, float] = {}

    for q in questions_on_page:
        q_num = q.get("question_number", 0)
        if not q_num:
            continue

        # The question number as its own word: "7." / "7)" or glued to the
        # first word of the stem ("7.Savol"). Note "12." does NOT match a
        # search for "1." — startswith compares the full "1." prefix.
        prefixes = (f"{q_num}.", f"{q_num})")
        matches = [
            w[1] for w in words
            if w[4] in prefixes or w[4].startswith(prefixes)
        ]
        if matches:
            y_positions[q_num] = min(matches)

    return y_positions


def _question_y_band(
    pdf_bytes: bytes,
    src_page: int,
    analysis_page: int,
    question_number: int,
    all_questions: list[dict],
    page_pdf_height: float,
    x_range: tuple[float, float] | None,
) -> tuple[float, float] | None:
    """
    Vertical band a question occupies: from its own number's Y position down
    to the next question's Y position within the SAME column (x_range).
    Returns None when the question's number can't be located in the PDF text —
    no band means no crop (never estimated).
    """
    questions_on_page = [
        q for q in all_questions if q.get("page_number") == analysis_page
    ]
    if not questions_on_page:
        return None

    y_positions = _get_question_y_positions(
        pdf_bytes, src_page, questions_on_page, x_range
    )
    this_q_y = y_positions.get(question_number)
    if this_q_y is None:
        return None

    next_q_y = page_pdf_height  # default: end of column
    for y in sorted(y_positions.values()):
        if y > this_q_y + 5:  # +5pt tolerance for same-line text
            next_q_y = y
            break
    return this_q_y, next_q_y


def _rect_in_xrange(rect: fitz.Rect, x_range: tuple[float, float] | None) -> bool:
    if x_range is None:
        return True
    cx = (rect.x0 + rect.x1) / 2
    return x_range[0] <= cx <= x_range[1]


def _find_image_rect_in_band(
    pdf_bytes: bytes,
    src_page: int,
    band: tuple[float, float],
    x_range: tuple[float, float] | None,
) -> fitz.Rect | None:
    """Embedded raster image whose center falls inside the question's band
    AND column. Largest wins; nearest-below allowed within 1.5x band span."""
    rects = [
        r for r in _get_image_rects_on_page(pdf_bytes, src_page)
        if _rect_in_xrange(r, x_range)
    ]
    if not rects:
        return None

    y_top, y_bottom = band
    in_band = [r for r in rects if y_top <= (r.y0 + r.y1) / 2 < y_bottom]
    if in_band:
        return max(in_band, key=lambda r: r.width * r.height)

    below = [(r, (r.y0 + r.y1) / 2) for r in rects if (r.y0 + r.y1) / 2 >= y_top]
    if below:
        closest = min(below, key=lambda x: x[1])
        if closest[1] - y_top < (y_bottom - y_top) * 1.5:
            return closest[0]
    return None


def _expand_with_labels(
    union: fitz.Rect,
    words: list,
    x_range: tuple[float, float] | None,
    band: tuple[float, float],
) -> fitz.Rect:
    """
    Grow a vector-figure rect to include the figure's OWN text labels — the
    numbers/letters printed directly above/below the strokes (e.g. a number
    line's "9,5 birlik", "B", "-3", "0", "2,5", "A"). Without this we'd crop a
    bare axis with no numbers, which is useless.

    Only text spans within (near) the figure's x-range and within a small
    vertical window above/below it are pulled in. Option markers ("A)".."E)")
    and question numbers are excluded, and the window never reaches into the
    options block below.
    """
    y_top, y_bottom = band
    x0, x1 = x_range if x_range else (union.x0, union.x1)
    xlo, xhi = x0 - 12, x1 + 12
    margin = max(union.height * 1.6, 18.0)
    win_top = max(y_top, union.y0 - margin)
    win_bot = min(y_bottom, union.y1 + margin)

    # never expand into the options block: stop just above the first option
    # marker sitting below the figure
    opt_below = [
        w[1] for w in words
        if _OPTION_MARKER_RE.match(str(w[4]).strip()) and (w[1] + w[3]) / 2 > union.y1
    ]
    if opt_below:
        win_bot = min(win_bot, min(opt_below) - 2)

    fig = fitz.Rect(union)
    for w in words:
        wx0, wy0, wx1, wy1 = w[0], w[1], w[2], w[3]
        token = str(w[4]).strip()
        cx, cy = (wx0 + wx1) / 2, (wy0 + wy1) / 2
        if not (xlo <= cx <= xhi) or not (win_top <= cy <= win_bot):
            continue
        if _OPTION_MARKER_RE.match(token) or _QNUM_WORD_RE.match(token):
            continue
        fig |= fitz.Rect(wx0, wy0, wx1, wy1)
    return fig


def _find_drawing_figure_rect(
    pdf_bytes: bytes,
    src_page: int,
    band: tuple[float, float],
    x_range: tuple[float, float] | None,
) -> fitz.Rect | None:
    """
    Vector diagram detection: union of drawing bboxes (lines, curves, shapes)
    inside the question's band and column, then EXPANDED to include the
    figure's own text labels. Accepts wide-and-flat figures (number lines,
    rulers, timelines, bracket diagrams); rejects only tiny specks and
    hairline rules. If nothing figure-like is found the caller attaches
    NOTHING rather than a garbage region.
    """
    y_top, y_bottom = band
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[src_page - 1]
        drawings = page.get_drawings()
        words = page.get_text("words")
        doc.close()
    except Exception as e:
        logger.warning("get_drawings_failed", page=src_page, error=str(e))
        return None

    # Accumulate the bounding box by explicit min/max — fitz's `|` treats a
    # zero-area rect (a thin horizontal axis is height 0 in vector terms) as
    # "empty" and drops it, which would collapse a number line to nothing.
    x0 = y0 = float("inf")
    x1 = y1 = float("-inf")
    found = False
    for d in drawings:
        r = d.get("rect")
        if r is None:
            continue
        cy = (r.y0 + r.y1) / 2
        if not (y_top <= cy < y_bottom) or not _rect_in_xrange(r, x_range):
            continue
        x0, y0 = min(x0, r.x0), min(y0, r.y0)
        x1, y1 = max(x1, r.x1), max(y1, r.y1)
        found = True

    if not found:
        return None
    union = fitz.Rect(x0, y0, x1, y1)
    # ASPECT/AREA-aware acceptance: a number line/ruler/timeline is legitimately
    # WIDE and FLAT. Reject only genuine non-figures:
    if union.width < 40:                       # too narrow to be a figure
        return None
    if union.height < 5:                       # a hairline rule / underline
        return None
    if union.width * union.height < 200:       # a tiny speck
        return None
    if union.height > (y_bottom - y_top) * 1.05:  # spills the whole band
        return None

    fig = _expand_with_labels(union, words, x_range, band)
    # small padding so the topmost/bottommost labels aren't clipped
    top = fig.y0 - 4
    # BUG: the top padding could dip into the stem line just above the figure,
    # baking a garbled "toping." sliver into the crop. Clamp the TOP (only) to
    # stay below the nearest text line above the figure. Bottom/sides keep their
    # padding so the number-line labels (9,5 birlik / B / A ...) are never lost.
    xa, xb = x_range if x_range else (fig.x0, fig.x1)
    above = [
        w[3] for w in words
        if w[3] <= fig.y0 and xa - 12 <= (w[0] + w[2]) / 2 <= xb + 12
    ]
    if above:
        top = max(top, max(above) + 1)
    return fitz.Rect(fig.x0 - 4, top, fig.x1 + 4, fig.y1 + 4)


# ── Crop sanity check (FIX 1: garbage detector) ──────────────────────────────

# Option marker word: "A)" .. "E)" (Latin or Cyrillic), possibly glued ("A)2")
_OPTION_MARKER_RE = re.compile(r'^[A-EА-Е]\)')
# Question-number word: "12." / "12)" but NOT decimals like "0.5"
_QNUM_WORD_RE = re.compile(r'^(\d{1,2})[.)](?!\d)')

MAX_CROP_PAGE_RATIO = 0.35  # a figure never legitimately covers >35% of a page


def _rect_is_garbage(
    pdf_bytes: bytes,
    src_page: int,
    rect: fitz.Rect,
    own_qnum: int,
    page_area: float,
) -> str | None:
    """
    Decide whether a candidate figure rect is actually a chunk of OTHER
    questions (the garbage-crop bug). Returns the rejection reason, or None
    if the rect looks like a clean figure.

    Signals:
    - covers more than MAX_CROP_PAGE_RATIO of the page
    - contains 2+ option markers ("A)".."E)") — schemes don't have options
    - contains another question's number word ("39." / "39)")
    """
    if page_area > 0 and (rect.width * rect.height) / page_area > MAX_CROP_PAGE_RATIO:
        return "too_large"

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        words = doc[src_page - 1].get_text("words", clip=rect)
        doc.close()
    except Exception as e:
        logger.warning("crop_sanity_inspect_failed", page=src_page, error=str(e))
        return None  # can't inspect — don't block the attach

    option_markers = 0
    for w in words:
        token = str(w[4]).strip()
        if _OPTION_MARKER_RE.match(token):
            option_markers += 1
        m = _QNUM_WORD_RE.match(token)
        if m and int(m.group(1)) != own_qnum:
            return "contains_question_number"
    if option_markers >= 2:
        return "contains_option_markers"
    return None


def _crop_echoes_stem(
    pdf_bytes: bytes,
    src_page: int,
    rect: fitz.Rect,
    stem_text: str | None,
    threshold: float = 0.6,
) -> bool:
    """
    FIX 6(b): text-echo detector. If most of the question's own stem tokens
    appear inside the crop, the crop is a screenshot of the question text,
    not a figure — reject it. Uses PDF words (no OCR needed).
    """
    if not stem_text:
        return False
    stem_tokens = set(re.findall(r'[^\W_]{3,}', stem_text.lower()))
    if len(stem_tokens) < 4:
        return False  # too little signal to judge
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        words = doc[src_page - 1].get_text("words", clip=rect)
        doc.close()
    except Exception:
        return False
    crop_tokens = {str(w[4]).lower().strip(".,;:!?()") for w in words}
    overlap = len(stem_tokens & crop_tokens) / len(stem_tokens)
    return overlap >= threshold


def restore_list_markers(
    questions: list[dict],
    pdf_bytes: bytes,
    col_map: dict[int, dict] | None,
) -> None:
    """
    ISSUE 1 (second half): after the question's own number was stripped from
    the stem, check the SOURCE words — if the token right after the number
    word is a list marker like "1)" and the stem doesn't start with it,
    restore it (the extractor sometimes swallows the marker together with
    the number). Best-effort, PDF-only, every restore logged.
    """
    if not pdf_bytes:
        return
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        sizes = {i + 1: (p.rect.width, p.rect.height) for i, p in enumerate(doc)}
        words_cache: dict[int, list] = {}

        for q in questions:
            page = q.get("page_number")
            mapping = (col_map or {}).get(page) if page else None
            if not mapping or mapping["src_page"] not in sizes:
                continue
            src = mapping["src_page"]
            if src not in words_cache:
                words_cache[src] = doc[src - 1].get_text("words")
            pdf_w = sizes[src][0]
            x0, x1 = mapping["x0"] * pdf_w, mapping["x1"] * pdf_w
            col_words = [
                w for w in words_cache[src] if x0 <= (w[0] + w[2]) / 2 <= x1
            ]
            n = q.get("question_number")
            nw = next(
                (w for w in col_words
                 if re.match(rf'^{n}[.)](?!\d)', str(w[4]))),
                None,
            )
            if nw is None:
                continue
            # first word to the right on the same line
            line = sorted(
                [w for w in col_words if abs(w[1] - nw[1]) < 3 and w[0] > nw[0]],
                key=lambda w: w[0],
            )
            if not line:
                continue
            mm = re.match(r'^(\d+[.)])(?!\d)', str(line[0][4]))
            if not mm:
                continue
            marker = mm.group(1)
            stem = q.get("question_text") or ""
            if not stem.lstrip().startswith(marker):
                q["question_text"] = f"{marker} {stem}"
                logger.info(
                    "list_marker_restored",
                    question=n, marker=marker,
                )
        doc.close()
    except Exception as e:
        logger.warning("restore_list_markers_failed", error=str(e))


def save_debug_crops(
    questions: list[dict],
    numbers: list[int],
    pdf_bytes: bytes | None,
    col_map: dict[int, dict] | None,
    src_pages: list[PageImage] | None,
    limit: int = 10,
) -> dict[int, str]:
    """
    ISSUE 3: for flagged/suspicious questions, save an UNfiltered crop of the
    question's source band so a human can eyeball the transcription against
    the original. Debug artifacts only — never attached to the PDF.
    """
    if not pdf_bytes or not src_pages:
        return {}
    src_lookup = {p.page_number: p.image for p in src_pages}
    saved: dict[int, str] = {}
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        sizes = {i + 1: (p.rect.width, p.rect.height) for i, p in enumerate(doc)}
        doc.close()
    except Exception:
        return {}

    for q in questions:
        n = q.get("question_number")
        if n not in numbers or len(saved) >= limit:
            continue
        page = q.get("page_number")
        mapping = (col_map or {}).get(page) if page else None
        if not mapping or mapping["src_page"] not in src_lookup \
                or mapping["src_page"] not in sizes:
            continue
        src = mapping["src_page"]
        pdf_w, pdf_h = sizes[src]
        x_range = (mapping["x0"] * pdf_w, mapping["x1"] * pdf_w)
        band = _question_y_band(pdf_bytes, src, page, n, questions, pdf_h, x_range)
        if not band:
            continue
        rect = fitz.Rect(x_range[0], max(0, band[0] - 4), x_range[1], band[1] + 4)
        try:
            img = src_lookup[src]
            sx, sy = img.size[0] / pdf_w, img.size[1] / pdf_h
            crop = img.crop((
                int(rect.x0 * sx), int(rect.y0 * sy),
                int(rect.x1 * sx), int(rect.y1 * sy),
            ))
            path = _ensure_image_dir() / f"debug_q{n}_p{src}_{uuid.uuid4().hex[:6]}.png"
            crop.save(str(path), format="PNG")
            saved[n] = str(path)
            logger.info("debug_crop_saved", question=n, path=str(path))
        except Exception as e:
            logger.warning("debug_crop_failed", question=n, error=str(e))
    return saved


def recrop_scheme_region(
    pdf_bytes: bytes,
    src_page: int,
    src_image: Image.Image,
    page_pdf_size: tuple[float, float],
    x_range: tuple[float, float] | None,
    q_num: int,
    analysis_page: int,
    all_questions: list[dict],
    stem_text: str | None = None,
) -> str | None:
    """
    FIX 3(a): geometric re-crop — the region between the question's first
    stem line and the start of its own options, within its column. Runs the
    garbage detector AND the text-echo detector on the result; returns a
    saved path or None.
    """
    band = _question_y_band(
        pdf_bytes, src_page, analysis_page, q_num,
        all_questions, page_pdf_size[1], x_range,
    )
    if not band:
        return None
    y_top, y_bottom = band

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        words = doc[src_page - 1].get_text("words")
        doc.close()
    except Exception:
        return None
    if x_range is not None:
        words = [w for w in words if x_range[0] <= (w[0] + w[2]) / 2 <= x_range[1]]

    # Trim the first stem line (words sharing the question number's baseline)
    first_line = [w for w in words if abs(w[1] - y_top) < 3]
    top = max((w[3] for w in first_line), default=y_top) + 2

    # Stop at the question's own first option marker
    option_ys = [
        w[1] for w in words
        if top < w[1] < y_bottom and _OPTION_MARKER_RE.match(str(w[4]).strip())
    ]
    bottom = min(option_ys) if option_ys else y_bottom

    if bottom - top < 15:
        return None

    x0 = x_range[0] + 2 if x_range else 0
    x1 = x_range[1] - 2 if x_range else page_pdf_size[0]
    rect = fitz.Rect(x0, top, x1, bottom)

    page_area = page_pdf_size[0] * page_pdf_size[1]
    reason = _rect_is_garbage(pdf_bytes, src_page, rect, q_num, page_area)
    if reason:
        logger.info("recrop_rejected", question=q_num, reason=reason)
        return None
    # FIX 6(b): a re-crop dominated by the question's own text is a
    # screenshot of the stem, not a figure.
    if _crop_echoes_stem(pdf_bytes, src_page, rect, stem_text):
        logger.info("recrop_rejected", question=q_num, reason="text_echo")
        return None

    return crop_and_save_image(
        page_image=src_image,
        rect_pdf=rect,
        page_pdf_size=page_pdf_size,
        question_number=q_num,
        page_number=src_page,
        padding_px=12,
        pdf_bytes=pdf_bytes,
    )


# ── Crop & save ───────────────────────────────────────────────────────────────

def crop_and_save_image(
    page_image: Image.Image,
    rect_pdf: fitz.Rect,
    page_pdf_size: tuple[float, float],
    question_number: int,
    page_number: int,
    padding_px: int = 8,
    pdf_bytes: bytes | None = None,
    top_limit_pt: float | None = None,
) -> str | None:
    """
    Crop a figure region and save it as PNG, returning the file path.

    When pdf_bytes is available the region is RE-RENDERED straight from the
    PDF at CROP_DPI via get_pixmap(clip=...) — sharp output regardless of the
    200-DPI page render. Falls back to cropping the PIL page render.

    top_limit_pt: if set, the crop's TOP edge is never padded ABOVE this y
    (PDF points). Used for a figure clamped just under a stem line, so the
    symmetric padding can't pull the stem's descenders back into the crop.
    Bottom and sides keep their padding.
    """
    if pdf_bytes is not None:
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            page = doc[page_number - 1]
            pad_pt = padding_px * 72.0 / DPI
            top = max(0.0, rect_pdf.y0 - pad_pt)
            if top_limit_pt is not None:
                top = max(top, top_limit_pt)
            clip = fitz.Rect(
                max(0, rect_pdf.x0 - pad_pt),
                top,
                min(page.rect.x1, rect_pdf.x1 + pad_pt),
                min(page.rect.y1, rect_pdf.y1 + pad_pt),
            )
            zoom = CROP_DPI / 72.0
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=clip, alpha=False)
            doc.close()
            if pix.width < 40 or pix.height < 40:
                logger.warning("crop_too_small", q=question_number, page=page_number)
                return None
            save_dir = _ensure_image_dir()
            filename = f"q{question_number}_p{page_number}_{uuid.uuid4().hex[:8]}.png"
            save_path = save_dir / filename
            pix.save(str(save_path))
            logger.info(
                "crop_saved_hidpi", question=question_number, page=page_number,
                size=f"{pix.width}x{pix.height}", path=str(save_path),
            )
            return str(save_path)
        except Exception as e:
            logger.warning(
                "hidpi_crop_failed", question=question_number, error=str(e)
            )
            # fall through to the PIL path

    try:
        pdf_w, pdf_h = page_pdf_size
        img_w, img_h = page_image.size

        scale_x = img_w / pdf_w
        scale_y = img_h / pdf_h

        left   = max(0, int(rect_pdf.x0 * scale_x) - padding_px)
        top    = max(0, int(rect_pdf.y0 * scale_y) - padding_px)
        if top_limit_pt is not None:
            top = max(top, int(top_limit_pt * scale_y))
        right  = min(img_w, int(rect_pdf.x1 * scale_x) + padding_px)
        bottom = min(img_h, int(rect_pdf.y1 * scale_y) + padding_px)

        if right - left < 20 or bottom - top < 20:
            logger.warning("crop_too_small", q=question_number, page=page_number)
            return None

        cropped = page_image.crop((left, top, right, bottom))
        save_dir = _ensure_image_dir()
        filename = f"q{question_number}_p{page_number}_{uuid.uuid4().hex[:8]}.png"
        save_path = save_dir / filename
        cropped.save(str(save_path), format="PNG")

        logger.info("crop_saved", question=question_number, page=page_number,
                    size=f"{right-left}x{bottom-top}", path=str(save_path))
        return str(save_path)
    except Exception as e:
        logger.error("crop_failed", question=question_number, page=page_number, error=str(e))
        return None


# ── Public API ────────────────────────────────────────────────────────────────

def attach_images_to_questions(
    questions: list[dict],
    page_images: list[PageImage],
    pdf_bytes: bytes | None = None,
    col_map: dict[int, dict] | None = None,
    src_pages: list[PageImage] | None = None,
) -> list[dict]:
    """
    For every question with has_image=True, isolate its figure region, crop
    it from the SOURCE page render, and store the path in q["image_path"].

    A figure is attached only when confidently isolated (embedded raster
    rect, or a vector-drawing cluster, inside the question's own band and
    column). Otherwise image_path stays None — the PDF renders the
    description box instead. NEVER a band-of-the-page fallback: those crops
    contained other questions' content on multi-column pages.

    Args:
        questions:   Output of AIAnalyzer.extract_all_questions(); their
                     page_number values refer to page_images entries.
        page_images: Analysis pages (possibly column-split halves).
        pdf_bytes:   Original PDF bytes (enables figure isolation).
        col_map:     From split_two_column_pages: {analysis_page:
                     {"src_page", "x0", "x1"}}. None = identity.
        src_pages:   Original (unsplit) page renders for cropping.
                     None = page_images are the sources.
    """
    if col_map is None:
        col_map = {
            p.page_number: {"src_page": p.page_number, "x0": 0.0, "x1": 1.0}
            for p in page_images
        }
    if src_pages is None:
        src_pages = page_images
    src_lookup: dict[int, Image.Image] = {
        p.page_number: p.image for p in src_pages
    }

    # Pre-compute PDF page sizes if we have the PDF
    pdf_page_sizes: dict[int, tuple[float, float]] = {}
    if pdf_bytes:
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for i, page in enumerate(doc):
                r = page.rect
                pdf_page_sizes[i + 1] = (r.width, r.height)
            doc.close()
        except Exception as e:
            logger.warning("pdf_page_size_fail", error=str(e))

    for q in questions:
        if not q.get("has_image"):
            continue
        if q.get("image_path"):
            continue  # already assigned

        # FIX 6(c): the stem already carries the transformation chain as
        # text — any image would only echo it. Attach nothing.
        if "→" in (q.get("question_text") or ""):
            logger.info(
                "image_skipped_chain_in_stem",
                question=q.get("question_number"),
            )
            q["image_path"] = None
            continue

        page_num = q.get("page_number")
        q_num    = q.get("question_number", 0)
        mapping  = col_map.get(page_num) if page_num else None

        if not mapping or mapping["src_page"] not in src_lookup:
            logger.warning("no_page_image", question=q_num, page=page_num)
            q["image_path"] = None
            continue

        src_page = mapping["src_page"]
        path: str | None = None

        if pdf_bytes and src_page in pdf_page_sizes:
            try:
                pdf_w, pdf_h = pdf_page_sizes[src_page]
                x_range = (mapping["x0"] * pdf_w, mapping["x1"] * pdf_w)
                band = _question_y_band(
                    pdf_bytes=pdf_bytes,
                    src_page=src_page,
                    analysis_page=page_num,
                    question_number=q_num,
                    all_questions=questions,
                    page_pdf_height=pdf_h,
                    x_range=x_range,
                )
                if band:
                    # Strategy 1: embedded raster image in band+column
                    rect = _find_image_rect_in_band(
                        pdf_bytes, src_page, band, x_range
                    )
                    # Strategy 2: vector-drawing cluster (drawn diagrams)
                    from_drawing = False
                    if rect is None:
                        rect = _find_drawing_figure_rect(
                            pdf_bytes, src_page, band, x_range
                        )
                        from_drawing = rect is not None
                    # FIX 1: garbage detector — reject crops containing other
                    # questions' numbers/options or covering too much page.
                    if rect is not None:
                        reason = _rect_is_garbage(
                            pdf_bytes, src_page, rect, q_num,
                            pdf_w * pdf_h,
                        )
                        if reason is None and _crop_echoes_stem(
                            pdf_bytes, src_page, rect, q.get("question_text")
                        ):
                            reason = "text_echo"  # FIX 6(b)
                        if reason:
                            logger.warning(
                                "crop_rejected_garbage",
                                question=q_num, page=page_num, reason=reason,
                            )
                            rect = None
                    if rect is not None:
                        path = crop_and_save_image(
                            page_image=src_lookup[src_page],
                            rect_pdf=rect,
                            page_pdf_size=pdf_page_sizes[src_page],
                            question_number=q_num,
                            page_number=src_page,
                            pdf_bytes=pdf_bytes,
                            # a drawing figure is clamped just under the stem;
                            # don't let padding pull the stem back in
                            top_limit_pt=rect.y0 if from_drawing else None,
                        )
                    # FIX 3(a): rejected/missing figure → geometric re-crop
                    # of the stem→options region (garbage-checked again).
                    if not path:
                        path = recrop_scheme_region(
                            pdf_bytes=pdf_bytes,
                            src_page=src_page,
                            src_image=src_lookup[src_page],
                            page_pdf_size=pdf_page_sizes[src_page],
                            x_range=x_range,
                            q_num=q_num,
                            analysis_page=page_num,
                            all_questions=questions,
                            stem_text=q.get("question_text"),
                        )
            except Exception as e:
                logger.warning("precise_crop_fail", question=q_num, error=str(e))

        if not path:
            logger.info(
                "no_confident_figure",
                question=q_num, page=page_num,
                detail="attaching nothing; PDF will show the description box",
            )
        q["image_path"] = path

    return questions


# ── DOCX / image helpers ──────────────────────────────────────────────────────

# ── DOCX → page images ────────────────────────────────────────────────────────
# The old docx_to_images only harvested EMBEDDED pictures, so a normal typed
# test (text + options, no pasted images) produced zero pages and "no
# questions found". We now RENDER the document body (paragraphs, tables,
# inline images) to page images with Pillow + the bundled DejaVu font, then
# feed them into the SAME vision pipeline as PDFs/scans. No new system deps.

_DOCX_PAGE_W = 1654   # A4 width  @ 200 DPI (matches pdf_to_images DPI)
_DOCX_PAGE_H = 2339   # A4 height @ 200 DPI
_DOCX_MARGIN = 120
_DOCX_BODY_PT = 30
_DOCX_LINE_H = 42
_DOCX_TABLE_PAD = 12


def _docx_font(size: int, bold: bool = False):
    """Bundled DejaVu (Cyrillic/Uzbek/math coverage); Pillow default on miss."""
    fonts_dir = Path(__file__).resolve().parent.parent / "assets" / "fonts"
    name = "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"
    try:
        return ImageFont.truetype(str(fonts_dir / name), size)
    except Exception:
        return ImageFont.load_default()


def _iter_docx_blocks(parent):
    """Yield paragraphs and tables IN DOCUMENT ORDER (standard python-docx
    recipe) so questions, their options and any tables stay interleaved as
    the teacher wrote them."""
    if isinstance(parent, _DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc  # table cell
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield _DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield _DocxTable(child, parent)


def _wrap_text(draw, text: str, font, max_w: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.split("\n"):
        words = raw_line.split(" ")
        cur = ""
        for w in words:
            trial = f"{cur} {w}".strip()
            if draw.textlength(trial, font=font) <= max_w or not cur:
                cur = trial
            else:
                lines.append(cur)
                cur = w
        lines.append(cur)
    return lines or [""]


def _para_inline_images(doc: Document, para) -> list[Image.Image]:
    """Inline images embedded in a paragraph (best-effort)."""
    out: list[Image.Image] = []
    for blip in para._p.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        try:
            blob = doc.part.related_parts[rid].blob
            out.append(Image.open(io.BytesIO(blob)).convert("RGB"))
        except Exception as e:
            logger.warning("docx_inline_image_skip", error=str(e))
    return out


class _DocxCanvas:
    """A simple paginating page canvas."""

    def __init__(self):
        self.pages: list[Image.Image] = []
        self._new_page()

    def _new_page(self):
        self.img = Image.new("RGB", (_DOCX_PAGE_W, _DOCX_PAGE_H), "white")
        self.draw = ImageDraw.Draw(self.img)
        self.y = _DOCX_MARGIN
        self.pages.append(self.img)

    def _ensure(self, height: int):
        if self.y + height > _DOCX_PAGE_H - _DOCX_MARGIN:
            self._new_page()

    def text_block(self, text: str, font, bold_font=None):
        if not text.strip():
            self.y += _DOCX_LINE_H // 2
            return
        max_w = _DOCX_PAGE_W - 2 * _DOCX_MARGIN
        for line in _wrap_text(self.draw, text, font, max_w):
            self._ensure(_DOCX_LINE_H)
            self.draw.text((_DOCX_MARGIN, self.y), line, font=font, fill="black")
            self.y += _DOCX_LINE_H

    def image_block(self, pic: Image.Image):
        max_w = _DOCX_PAGE_W - 2 * _DOCX_MARGIN
        scale = min(1.0, max_w / pic.width)
        w, h = int(pic.width * scale), int(pic.height * scale)
        if h > _DOCX_PAGE_H - 2 * _DOCX_MARGIN:
            h2 = _DOCX_PAGE_H - 2 * _DOCX_MARGIN
            w = int(w * h2 / h)
            h = h2
        self._ensure(h + 20)
        self.img.paste(pic.resize((w, h)), (_DOCX_MARGIN, self.y))
        self.y += h + 20

    def table_block(self, table, font):
        max_w = _DOCX_PAGE_W - 2 * _DOCX_MARGIN
        rows = table.rows
        if not rows:
            return
        ncols = max(len(r.cells) for r in rows)
        col_w = max_w // max(1, ncols)
        for row in rows:
            cells = row.cells
            wrapped = [
                _wrap_text(self.draw, c.text, font, col_w - 2 * _DOCX_TABLE_PAD)
                for c in cells
            ]
            row_h = max(
                (len(w) * _DOCX_LINE_H + 2 * _DOCX_TABLE_PAD for w in wrapped),
                default=_DOCX_LINE_H,
            )
            self._ensure(row_h)
            x = _DOCX_MARGIN
            for ci in range(ncols):
                self.draw.rectangle(
                    [x, self.y, x + col_w, self.y + row_h], outline="black", width=2
                )
                ty = self.y + _DOCX_TABLE_PAD
                for line in (wrapped[ci] if ci < len(wrapped) else []):
                    self.draw.text((x + _DOCX_TABLE_PAD, ty), line, font=font, fill="black")
                    ty += _DOCX_LINE_H
                x += col_w
            self.y += row_h


def docx_to_images(docx_bytes: bytes) -> tuple[list[PageImage], str]:
    """Render a DOCX into page images (fed to the vision pipeline) + its text.

    Returns (pages, text). Falls back to a plain-text render if structured
    walking fails, so a malformed DOCX still produces at least one page.
    """
    doc = Document(io.BytesIO(docx_bytes))
    canvas = _DocxCanvas()
    body_font = _docx_font(_DOCX_BODY_PT)
    text_parts: list[str] = []

    try:
        for block in _iter_docx_blocks(doc):
            if isinstance(block, _DocxParagraph):
                if block.text.strip():
                    canvas.text_block(block.text, body_font)
                    text_parts.append(block.text)
                for pic in _para_inline_images(doc, block):
                    canvas.image_block(pic)
            elif isinstance(block, _DocxTable):
                canvas.table_block(block, body_font)
                for row in block.rows:
                    text_parts.append(" | ".join(c.text for c in row.cells))
    except Exception as e:
        logger.warning("docx_structured_render_failed", error=str(e))
        # Fallback: dump all paragraph text onto fresh pages.
        canvas = _DocxCanvas()
        for para in doc.paragraphs:
            if para.text.strip():
                canvas.text_block(para.text, body_font)
                text_parts.append(para.text)

    pages = [
        PageImage(page_number=i + 1, image=img)
        for i, img in enumerate(canvas.pages)
    ]
    text_content = "\n".join(text_parts)
    logger.info(
        "docx_rendered", pages=len(pages), text_chars=len(text_content)
    )
    return pages, text_content


def _save_docx_image(pic: Image.Image) -> str | None:
    """Save one decoded DOCX inline image as PNG in IMAGE_SAVE_DIR.

    The `docximg_` filename prefix is a deliberate marker: the PDF renderer
    reads it to size the image fit-to-column. A CROP_DPI crop has a known
    physical size (pixels / CROP_DPI inches); a DOCX-embedded image does NOT —
    it is arbitrary pixels — so the two must be sized differently.
    """
    try:
        save_dir = _ensure_image_dir()
        path = save_dir / f"docximg_{uuid.uuid4().hex[:10]}.png"
        pic.save(str(path), format="PNG")
        logger.info(
            "docx_image_saved", path=str(path), size=f"{pic.width}x{pic.height}"
        )
        return str(path)
    except Exception as e:
        logger.warning("docx_image_save_failed", error=str(e))
        return None


def docx_extract_ordered_images(docx_bytes: bytes) -> list[Image.Image]:
    """Every inline image in the DOCX, IN DOCUMENT (reading) ORDER.

    Kept SEPARATE from docx_to_images (whose 2-tuple return other callers rely
    on). Best-effort: a malformed document or an undecodable picture is
    skipped, never raised. The a:blip -> r:embed -> media lookup lives in
    _para_inline_images and is reused verbatim.
    """
    out: list[Image.Image] = []
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        logger.warning("docx_open_failed", error=str(e))
        return out
    try:
        for block in _iter_docx_blocks(doc):
            if isinstance(block, _DocxParagraph):
                out.extend(_para_inline_images(doc, block))
            elif isinstance(block, _DocxTable):
                # images inside table cells, cell-by-cell in reading order
                for row in block.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            out.extend(_para_inline_images(doc, para))
    except Exception as e:
        logger.warning("docx_ordered_images_failed", error=str(e))
    return out


def attach_docx_inline_images(questions: list[dict], docx_bytes: bytes) -> int:
    """Attach a DOCX's embedded images to the questions Gemini flagged.

    DOCX has no page geometry, so figures can't be cropped by rect like a PDF.
    Instead we pair the ordered inline images with the flagged questions IN
    QUESTION ORDER — but ONLY when the two counts match EXACTLY. A mismatch
    attaches nothing and the description box still renders, mirroring the PDF
    rule (principle: never attach a WRONG figure). Only fills questions that
    have has_image=True and no image_path yet. Returns the count attached.
    """
    images = docx_extract_ordered_images(docx_bytes)
    if not images:
        return 0
    flagged = sorted(
        (q for q in questions
         if q.get("has_image") and not q.get("image_path")),
        key=lambda q: q.get("question_number", 0),
    )
    if len(flagged) != len(images):
        logger.info(
            "docx_image_count_mismatch",
            images=len(images), flagged=len(flagged),
        )
        return 0
    attached = 0
    for q, pic in zip(flagged, images):
        path = _save_docx_image(pic)
        if path:
            q["image_path"] = path
            attached += 1
    logger.info("docx_images_attached", count=attached)
    return attached


def docx_extract_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def image_to_pages(image_bytes: bytes) -> list[PageImage]:
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    return [PageImage(page_number=1, image=img)]


def preprocess_image(img: Image.Image) -> Image.Image:
    import numpy as np
    import cv2

    arr = np.array(img)
    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)

    coords = np.column_stack(np.where(gray < 200))
    if coords.size > 0:
        angle = cv2.minAreaRect(coords)[-1]
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle
        if abs(angle) > 0.5:
            h, w = gray.shape
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            gray = cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC,
                                  borderMode=cv2.BORDER_REPLICATE)

    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(gray)
    denoised = cv2.fastNlMeansDenoising(enhanced, h=10)
    rgb = cv2.cvtColor(denoised, cv2.COLOR_GRAY2RGB)
    return Image.fromarray(rgb)


def image_to_bytes(img: Image.Image, fmt: str = "PNG") -> bytes:
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()